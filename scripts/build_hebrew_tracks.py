from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "workbooks-source"
PDF_DIR = ROOT / "public" / "downloads"

GOLD = "B68A32"
NAVY = "21364A"
CREAM = "F8F2E7"
PALE_BLUE = "EAF2F6"
PALE_GOLD = "F5E8C8"
INK = "25313A"
MUTED = "5C6870"


FOUNDATION = [
    {
        "title": "Quiet Doorways: Aleph and Ayin",
        "focus": "א  ע   ַ   ָ",
        "idea": "Aleph and Ayin are quiet letter doorways in the pronunciation we are learning. A vowel mark supplies the voice.",
        "teacher": "Point to the letter first, then look underneath. The letter is quiet; the mark tells our mouth to say ah.",
        "notice": [
            "Patach is our stick vowel: a short line under the letter.",
            "Open your mouth as if the doctor says, ‘Say ahhh.’",
            "Kamatz is the stick vowel with a handle. In today’s words, it also says ah.",
        ],
        "examples": ["אַ   עַ", "אָ   עָ", "אַ  אָ  עַ  עָ"],
        "practice": [
            "Circle every stick vowel. Box every stick vowel with a handle.",
            "Point from right to left and say each sound: אַ  אָ  עַ  עָ",
            "Trace: א  א  א     ע  ע  ע",
            "Draw a short stick under א. Then add a handle to make Kamatz.",
        ],
        "check": "I can find Aleph, Ayin, Patach, and Kamatz.",
        "havari": "How does the mark help a quiet letter speak?",
        "note": "Pronunciation note: Ayin is heard as a throat sound in some Hebrew traditions. Kamatz can sometimes sound like oh; this lesson uses common beginner ah examples.",
    },
    {
        "title": "Ball or Void: Bet and Vet",
        "focus": "בּ  ב   ְ",
        "idea": "Bet and Vet share one body. A ball in the center cues b; a void in the center cues v.",
        "teacher": "Look inside before you speak: ball means Bet, b. Void means Vet, v.",
        "notice": [
            "Bet has a ball in the center: בּ = b.",
            "Vet has a void where the ball would be: ב = v.",
            "Sheva is two tiny footsteps. Sometimes they move with a very short helper sound; sometimes they rest and are silent.",
        ],
        "examples": ["בַּ  בָּ", "בַ  בָ", "בְּרָכָה   מַלְכָּה"],
        "practice": [
            "Circle the ball: ב  בּ  ב  בּ",
            "Underline the void: בּ  ב  בּ  ב",
            "Point and say: בַּ  בָּ  בַ  בָ",
            "For each Sheva, let the teacher model whether the footsteps move or rest.",
        ],
        "check": "I look inside the letter before choosing b or v.",
        "havari": "Are the Sheva footsteps moving or resting in this word?",
        "note": "Sheva note: do not teach one fixed sound. At a word’s beginning it often supports a very short sound; elsewhere it is often silent. Model the whole word before asking the learner to decide.",
    },
]


TRACKS = {
    "Aleph": [
        ("Mem Has Two Homes", "מ  ם", "Mem says m. Final Mem is used only at the end of a word.", ["Find the open Mem: מ", "Find the closed Final Mem: ם", "Trace both shapes from right to left."], ["מַ  מָ", "מַם"]),
        ("Shin and Sin: The Dot Moves", "שׁ  שׂ", "The dot on the right cues sh; the dot on the left cues s.", ["Touch the dot before saying the sound.", "Circle Shin; underline Sin.", "Build שַׁ and שָׂ."], ["שַׁ  שָׁ", "שַׂ  שָׂ"]),
        ("Lamed Stands Tall", "ל", "Lamed rises above the writing line and says l.", ["Find the tallest letter.", "Trace ל three times.", "Blend לַ and לָ."], ["לַ  לָ", "לַב"]),
        ("Vav and Yod: Small Shapes Matter", "ו  י", "Vav and Yod are small, distinct shapes that help build many words.", ["Compare height and length.", "Circle Vav; box Yod.", "Read from right to left."], ["וַ  וָ", "יַ  יָ"]),
        ("Hey and Chet: Find the Opening", "ה  ח", "Hey has an opening; Chet is closed across the top and sides.", ["Point to the opening in ה.", "Circle Chet.", "Say h gently; let the teacher model ch."], ["הַ  הָ", "חַ  חָ"]),
        ("Tav Finishes the Line", "ת", "Tav says t and often appears at the end of words.", ["Trace Tav.", "Find ת at the end.", "Blend תַ and תָ."], ["תַ  תָ", "בַּת"]),
        ("The Dot Vowel: Chirik", "ִ", "A single dot under a letter is Chirik and usually cues ee.", ["Find the dot under the letter.", "Contrast ah and ee.", "Blend slowly."], ["בִּ  מִ  לִ", "בַּ  בִּ"]),
        ("Three-Dot Steps: Segol", "ֶ", "Three dots under a letter are Segol and cue eh.", ["Count three dots.", "Say eh without adding an extra syllable.", "Compare eh with ah."], ["בֶּ  מֶ  לֶ", "בַּ  בֶּ"]),
        ("Build a Sound, Then a Word", "אוֹר  בַּת", "Readers blend one marked letter at a time, then say the whole word.", ["Point to each sound.", "Sweep under the whole word.", "Say it again smoothly."], ["בַּת", "אוֹר"]),
        ("Aleph Track Celebration", "א  ע  בּ  ב  מ  ם  שׁ  ל", "A careful reader looks, points, sounds, and reflects.", ["Name five letters.", "Find two vowel marks.", "Choose one word to read."], ["בַּת", "אוֹר", "שָׁלוֹם"]),
    ],
    "Bet": [
        ("Gimel and Dalet", "ג  ד", "Gimel says g; Dalet says d. Notice the different roof and foot.", ["Trace each shape.", "Circle Gimel.", "Blend with ah."], ["גַּ  גָּ", "דַּ  דָּ"]),
        ("Kaf and Khaf", "כּ  כ  ך", "A dot cues k; a void cues the breathy kh sound. Final Khaf ends a word.", ["Look inside first.", "Find Final Khaf.", "Blend with known vowels."], ["כַּ  כָּ", "כַ  כָ"]),
        ("Peh and Feh", "פּ  פ  ף", "A dot cues p; a void cues f. Final Feh ends a word.", ["Circle every ball.", "Underline every void.", "Find the final form."], ["פַּ  פָּ", "פַ  פָ"]),
        ("Tsadi and Final Tsadi", "צ  ץ", "Tsadi says ts and changes shape at the end of a word.", ["Trace both forms.", "Sort middle and final shapes.", "Blend slowly."], ["צַ  צָ", "אֶרֶץ"]),
        ("Final-Form Family", "ך  ם  ן  ף  ץ", "Five letters change shape only at the end of a word.", ["Name each family pair.", "Circle final forms.", "Place a final form only in the ending box."], ["מ / ם", "כ / ך", "פ / ף"]),
        ("Sheva in Whole Words", "ְ", "Sheva can be moving or resting; the whole word tells us how to read it.", ["Mark moving footsteps in blue.", "Mark resting footsteps in gray.", "Listen, echo, then read."], ["בְּרָכָה", "מַלְכָּה"]),
        ("Open and Closed Syllables", "בַּ  בַּת", "A syllable may end in a vowel sound or close with a consonant.", ["Tap each syllable.", "Draw a box around the closing consonant.", "Blend without rushing."], ["בַּ", "בַּת", "מֶלֶךְ"]),
        ("Read אור: Light", "אוֹר", "Known letters and vowels join into a meaningful word: or, light.", ["Point to each letter.", "Sweep and read.", "Connect the word to meaning."], ["אוֹר", "הָאוֹר"]),
        ("Read בית and שלום", "בַּיִת  שָׁלוֹם", "Readers carry sound into meaning, one word at a time.", ["Mark the syllables.", "Find the ball in Bet.", "Explain each meaning."], ["בַּיִת", "שָׁלוֹם"]),
        ("Bet Track Fluency", "אוֹר  בַּיִת  שָׁלוֹם", "Fluency is accurate, gentle reading - not racing.", ["Read once accurately.", "Read again smoothly.", "Tell one meaning."], ["אוֹר", "בַּיִת", "שָׁלוֹם"]),
    ],
    "Gimel": [
        ("Decode Before You Guess", "מִלָּה", "A strong reader notices letters and marks before using context.", ["Cover the picture.", "Point and decode.", "Check meaning afterward."], ["אוֹר", "בַּיִת", "שָׁלוֹם"]),
        ("See, Read, Mean", "רֹאִים • קוֹרְאִים • מְבִינִים", "Gimel Track connects accurate seeing, fluent reading, and meaning.", ["Name what you see.", "Read the word.", "Explain or sketch the meaning."], ["מַיִם", "יוֹם", "לַיְלָה"]),
        ("Build a Vocabulary Card", "מִלִּים", "A useful word card records Hebrew, pronunciation, meaning, and an example.", ["Copy exactly.", "Mark syllables.", "Use the word in a short phrase."], ["סֵפֶר", "מוֹרֶה", "תַּלְמִיד"]),
        ("Roots Are Meaning Families", "ש־ל־ם", "A root is a recurring letter family that can connect related meanings.", ["Circle the root letters.", "Compare related forms.", "State the shared meaning idea."], ["שָׁלוֹם", "שָׁלֵם"]),
        ("The Prefix ה: The", "הַ", "At the beginning of many words, ה can mark the definite article, the.", ["Box the prefix.", "Read the base word.", "Read the whole word."], ["אוֹר / הָאוֹר", "בַּיִת / הַבַּיִת"]),
        ("The Prefix ו: And", "וְ", "Vav can join words with the meaning and; its vowel changes in some contexts.", ["Circle the connector.", "Read each word separately.", "Read the phrase smoothly."], ["אוֹר וְשָׁלוֹם", "יוֹם וָלַיְלָה"]),
        ("Meaning from Context", "הֶקְשֵׁר", "Nearby words and ideas can confirm meaning after decoding.", ["Decode the target word.", "Read around it.", "Choose the meaning that fits."], ["בַּיִת", "סֵפֶר", "דֶּרֶךְ"]),
        ("Read a Short Source Line", "קוֹל אוֹר", "A source line deserves slow reading, respectful questions, and teacher guidance.", ["Locate familiar words.", "Mark what is new.", "Ask one honest question."], ["אוֹר", "דֶּרֶךְ שָׁלוֹם"]),
        ("Use a Dictionary Entry", "מִלּוֹן", "A dictionary confirms spelling, pronunciation, part of speech, and meaning.", ["Find the headword.", "Check the root or form.", "Record only the meaning that fits."], ["אוֹר", "שָׁלוֹם"]),
        ("Gimel Track Teach-Back", "לִרְאוֹת • לִקְרוֹא • לְהָבִין", "Understanding grows when a learner can explain the reading path clearly.", ["Choose one word.", "Model decoding.", "Explain its meaning and one useful detail."], ["מַיִם", "בַּיִת", "שָׁלוֹם"]),
    ],
}


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def font(run, name="Aptos", size=11, bold=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Arial")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=11, bold=False, color=INK, align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text); font(r, size=size, bold=bold, color=color)
    return p


def add_label(doc, label, text, fill=PALE_BLUE):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    c = t.cell(0, 0); c.width = Inches(6.5); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(c); shade(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    a = p.add_run(label.upper() + "  "); font(a, size=9, bold=True, color=GOLD)
    b = p.add_run(text); font(b, size=11, color=INK)
    add_text(doc, "", after=2)


def add_footer(section, track):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Ivrit HaOr • {track} Track • LuminaNexus Foundation"); font(r, size=8, color=MUTED)


def configure(doc, track):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    add_footer(sec, track)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (("Heading 1",16,NAVY,18,10),("Heading 2",13,NAVY,14,7),("Heading 3",12,NAVY,10,5)):
        s=doc.styles[style_name]; s.font.name="Aptos Display"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True


def page_break(doc):
    doc.add_page_break()


def cover(doc, track, lesson_count):
    add_text(doc, "IVRIT HAOR", 10, True, GOLD, WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_text(doc, f"{track} Track", 32, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "Student Workbook", 18, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_text(doc, "See carefully. Say gently. Build courage.", 14, False, INK, WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_label(doc, "Student", "____________________________________________", CREAM)
    add_label(doc, "Hebrew name", "____________________________________________", CREAM)
    add_text(doc, f"{lesson_count} lessons • fixed two-page lesson format • printable", 10, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, before=30)
    page_break(doc)
    add_text(doc, "How this workbook works", 22, True, NAVY, keep=True)
    add_text(doc, "Every lesson uses the same learning rhythm. Familiar structure leaves more attention for Hebrew.", 12, after=14)
    for title, body in [
        ("1. LOOK", "Name the letter or mark before making a sound."),
        ("2. NOTICE", "Use one memorable visual cue, then check the real feature."),
        ("3. SAY", "Listen to a teacher model, echo, and keep the sound brief."),
        ("4. BUILD", "Join a letter and vowel, then blend into a word."),
        ("5. CHECK", "Show what is secure and name what still needs practice."),
        ("6. ASK", "Use the Havari question to make learning conversational."),
    ]: add_label(doc, title, body, PALE_BLUE if title[0] in "135" else CREAM)
    add_text(doc, "Teacher accuracy promise", 13, True, NAVY, before=10)
    add_text(doc, "Mnemonics are doors, not definitions. Teacher notes preserve pronunciation differences and exceptions without overloading the learner.", 10, color=MUTED)


def lesson_page(doc, n, item, track):
    title=item["title"] if isinstance(item,dict) else item[0]
    focus=item["focus"] if isinstance(item,dict) else item[1]
    idea=item["idea"] if isinstance(item,dict) else item[2]
    notice=item["notice"] if isinstance(item,dict) else item[3]
    examples=item["examples"] if isinstance(item,dict) else item[4]
    teacher=item.get("teacher", f"Look first. Name what you notice. Then say or read {focus} slowly.") if isinstance(item,dict) else f"Look first. Name what you notice. Then say or read {focus} slowly."
    check=item.get("check", f"I can recognize and use {focus}.") if isinstance(item,dict) else f"I can recognize and use {focus}."
    havari=item.get("havari", "What helped you read carefully today?") if isinstance(item,dict) else "What helped you read carefully today?"
    note=item.get("note", "Model the complete sound or word before independent reading. Accept careful approximations, then recast accurately.") if isinstance(item,dict) else "Model the complete sound or word before independent reading. Accept careful approximations, then recast accurately."

    page_break(doc)
    add_text(doc, f"{track.upper()} TRACK  •  LESSON {n}", 9, True, GOLD, after=6)
    add_text(doc, title, 23, True, NAVY, after=10)
    add_text(doc, focus, 30, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, before=4, after=14)
    add_label(doc, "Big idea", idea, PALE_GOLD)
    add_label(doc, "Teacher says", teacher, PALE_BLUE)
    add_text(doc, "Notice", 14, True, NAVY, before=8, after=6)
    for x in notice:
        p=add_text(doc, "• "+x, 11, after=5); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.18)
    add_text(doc, "See it together", 14, True, NAVY, before=8, after=6)
    for x in examples: add_text(doc, x, 19, True, INK, WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_label(doc, "Teacher note", note, CREAM)

    page_break(doc)
    add_text(doc, f"LESSON {n} PRACTICE", 9, True, GOLD, after=6)
    add_text(doc, title, 20, True, NAVY, after=12)
    add_text(doc, "Look • Point • Say • Write", 13, True, NAVY, after=8)
    for i,x in enumerate(notice,1):
        add_text(doc, f"{i}. {x}", 11, after=9)
        add_text(doc, "____________________________________________________________", 10, color="AEB8BE", after=9)
    add_text(doc, "Read or name", 13, True, NAVY, before=4, after=7)
    add_text(doc, "     ".join(examples), 18, True, INK, WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_label(doc, "Confidence check", "□ Not yet   □ With help   □ By myself\n"+check, PALE_BLUE)
    add_label(doc, "Havari question", havari, PALE_GOLD)
    add_text(doc, "My best noticing today:", 11, True, NAVY, before=8)
    add_text(doc, "____________________________________________________________", 10, color="AEB8BE")


def make_track(track):
    lessons=FOUNDATION + [dict(title=a,focus=b,idea=c,notice=d,examples=e) for a,b,c,d,e in TRACKS[track]]
    doc=Document(); configure(doc, track); cover(doc, track, len(lessons))
    for n,item in enumerate(lessons,1): lesson_page(doc,n,item,track)
    SOURCE_DIR.mkdir(parents=True,exist_ok=True)
    out=SOURCE_DIR/f"{track}_Track_Rewritten.docx"; doc.save(out)
    return out


if __name__ == "__main__":
    for name in TRACKS:
        print(make_track(name))
